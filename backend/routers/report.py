from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
import json, os, uuid
from datetime import datetime
from database import get_db
from auth_utils import get_current_user
import models

router = APIRouter()

REPORTS_DIR = "generated_reports"
os.makedirs(REPORTS_DIR, exist_ok=True)


def generate_pdf_report(query_text: str, analysis: dict, user_name: str) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    filename = f"{REPORTS_DIR}/legal_report_{uuid.uuid4().hex[:8]}.pdf"
    doc = SimpleDocTemplate(filename, pagesize=A4,
                            rightMargin=0.75*inch, leftMargin=0.75*inch,
                            topMargin=1*inch, bottomMargin=0.75*inch)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'],
                                  fontSize=18, textColor=colors.HexColor('#1a365d'),
                                  alignment=TA_CENTER, spaceAfter=6)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'],
                                    fontSize=13, textColor=colors.HexColor('#2c5282'),
                                    spaceBefore=12, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                 fontSize=10, leading=14, spaceAfter=4)
    label_style = ParagraphStyle('Label', parent=styles['Normal'],
                                  fontSize=10, textColor=colors.HexColor('#718096'),
                                  spaceAfter=2)

    story = []

    # Header
    story.append(Paragraph("⚖ AI Legal Assistant", title_style))
    story.append(Paragraph("Legal Analysis Report", ParagraphStyle('sub', parent=styles['Normal'],
                             fontSize=12, alignment=TA_CENTER, textColor=colors.HexColor('#4a5568'))))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#2c5282')))
    story.append(Spacer(1, 12))

    # Meta
    meta_data = [
        ["Generated For:", user_name],
        ["Date:", datetime.now().strftime("%d %B %Y, %I:%M %p")],
        ["Legal Category:", analysis.get("legal_category", "N/A")],
    ]
    meta_table = Table(meta_data, colWidths=[1.8*inch, 4.5*inch])
    meta_table.setStyle(TableStyle([
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#718096')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 12))

    # Query
    story.append(Paragraph("Your Legal Query", heading_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 4))
    story.append(Paragraph(query_text, body_style))
    story.append(Spacer(1, 8))

    # Summary
    story.append(Paragraph("Legal Summary", heading_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 4))
    story.append(Paragraph(analysis.get("summary", ""), body_style))
    story.append(Spacer(1, 8))

    # IPC Sections
    story.append(Paragraph("Relevant Legal Sections", heading_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 4))

    for sec in analysis.get("relevant_sections", []):
        story.append(Paragraph(f"<b>{sec.get('section_number')} – {sec.get('title')}</b>", body_style))
        story.append(Paragraph(sec.get("description", ""), body_style))
        story.append(Paragraph(f"<b>Punishment:</b> {sec.get('punishment', '')}", body_style))
        if sec.get("fine"):
            story.append(Paragraph(f"<b>Fine:</b> {sec.get('fine')}", body_style))
        story.append(Spacer(1, 6))

    # Outcomes
    story.append(Paragraph("Possible Legal Outcomes", heading_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 4))
    for i, outcome in enumerate(analysis.get("possible_outcomes", []), 1):
        story.append(Paragraph(f"{i}. {outcome}", body_style))
    story.append(Spacer(1, 8))

    # Precautions
    story.append(Paragraph("Precautions", heading_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 4))
    for p in analysis.get("precautions", []):
        story.append(Paragraph(f"• {p}", body_style))
    story.append(Spacer(1, 8))

    # Recommended Actions
    story.append(Paragraph("Recommended Actions", heading_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 4))
    for a in analysis.get("recommended_actions", []):
        story.append(Paragraph(f"→ {a}", body_style))
    story.append(Spacer(1, 12))

    # Disclaimer
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 6))
    disclaimer = ("DISCLAIMER: This report is generated by an AI system for informational purposes only. "
                  "It does not constitute legal advice. Please consult a qualified legal professional "
                  "for advice specific to your situation.")
    story.append(Paragraph(disclaimer, ParagraphStyle('disclaimer', parent=styles['Normal'],
                             fontSize=8, textColor=colors.HexColor('#a0aec0'), alignment=TA_CENTER)))

    doc.build(story)
    return filename


def generate_docx_report(query_text: str, analysis: dict, user_name: str) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    filename = f"{REPORTS_DIR}/legal_report_{uuid.uuid4().hex[:8]}.docx"
    doc = Document()

    # Title
    title = doc.add_heading("AI Legal Assistant – Legal Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc.add_paragraph(f"Generated for: {user_name} | Date: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph(f"Legal Category: {analysis.get('legal_category', 'N/A')}")
    doc.add_paragraph("─" * 60)

    doc.add_heading("Your Legal Query", level=1)
    doc.add_paragraph(query_text)

    doc.add_heading("Legal Summary", level=1)
    doc.add_paragraph(analysis.get("summary", ""))

    doc.add_heading("Relevant Legal Sections", level=1)
    for sec in analysis.get("relevant_sections", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{sec.get('section_number')} – {sec.get('title')}")
        run.bold = True
        doc.add_paragraph(sec.get("description", ""))
        doc.add_paragraph(f"Punishment: {sec.get('punishment', '')}")

    doc.add_heading("Possible Legal Outcomes", level=1)
    for i, outcome in enumerate(analysis.get("possible_outcomes", []), 1):
        doc.add_paragraph(f"{i}. {outcome}")

    doc.add_heading("Precautions", level=1)
    for p in analysis.get("precautions", []):
        doc.add_paragraph(f"• {p}")

    doc.add_heading("Recommended Actions", level=1)
    for a in analysis.get("recommended_actions", []):
        doc.add_paragraph(f"→ {a}")

    doc.add_paragraph("─" * 60)
    doc.add_paragraph(
        "DISCLAIMER: This report is for informational purposes only and does not constitute legal advice. "
        "Please consult a qualified advocate for your specific situation."
    ).italic = True

    doc.save(filename)
    return filename


@router.get("/download/{query_id}/{format}")
def download_report(
    query_id: int,
    format: str,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    if format not in ["pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Format must be 'pdf' or 'docx'")

    q = db.query(models.UserQuery).filter(
        models.UserQuery.query_id == query_id,
        models.UserQuery.user_id == current_user.id
    ).first()
    if not q:
        raise HTTPException(status_code=404, detail="Query not found")

    analysis = json.loads(q.analysis_result) if q.analysis_result else {}

    try:
        if format == "pdf":
            filepath = generate_pdf_report(q.query_text, analysis, current_user.name)
            media_type = "application/pdf"
        else:
            filepath = generate_docx_report(q.query_text, analysis, current_user.name)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return FileResponse(
            filepath,
            media_type=media_type,
            filename=f"legal_report_{query_id}.{format}"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")
